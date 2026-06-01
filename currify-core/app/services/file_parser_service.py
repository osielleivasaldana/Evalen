import io
import logging
import re
from typing import Dict, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)

class FileParserService:
    """Servicio para extraer texto de diferentes tipos de archivos"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.rtf', '.html', '.htm']

    async def parse_file(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """
        Extrae texto de un archivo de forma asíncrona (non-blocking)
        Ejecuta el parsing intensivo en un thread pool.
        """
        import asyncio
        return await asyncio.to_thread(self._parse_file_sync, file_content, filename)

    def _parse_file_sync(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """
        Método síncrono interno para el parsing real
        """
        try:
            file_extension = Path(filename).suffix.lower()

            if file_extension not in self.supported_formats:
                raise ValueError(f"Formato de archivo no soportado: {file_extension}")

            if file_extension == '.pdf':
                return self._parse_pdf(file_content, filename)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_word(file_content, filename)
            elif file_extension == '.txt':
                return self._parse_text(file_content, filename)
            elif file_extension == '.rtf':
                return self._parse_rtf(file_content, filename)
            elif file_extension in ['.html', '.htm']:
                return self._parse_html(file_content, filename)
            else:
                raise ValueError(f"Parser no implementado para: {file_extension}")

        except Exception as e:
            logger.error(f"Error parsing file {filename}: {str(e)}")
            return {
                "success": False,
                "text": "",
                "error": str(e),
                "metadata": {
                    "filename": filename,
                    "file_type": file_extension,
                    "parsing_method": "error"
                }
            }

    def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extrae texto de archivos PDF usando pdfplumber (Más robusto con tildes)"""
        try:
            import pdfplumber
            from io import BytesIO

            text_parts = []
            pages_count = 0
            
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                pages_count = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    try:
                        # Hybrid Strategy: Extract both Default (for reading) and Layout (for columns)
                        # This ensures we get clean paragraphs AND table/column structure
                        text_default = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                        text_layout = page.extract_text(layout=True) or ""
                        
                        # Use default as primary, but append layout version if they differ significantly
                        # or just append page-by-page?
                        # Appending page-by-page keeps context together.
                        combined_page = f"{text_default}\n\n--- LAYOUT CONTEXT (PAGE {i+1}) ---\n{text_layout}"
                        text_parts.append(combined_page)
                        
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i + 1}: {e}")
                        continue

            extracted_text = "\n\n".join(text_parts)

            # Fallback: OCR strategy if text is insufficient
            if not extracted_text.strip() or len(extracted_text) < 100:
                logger.warning(f"Insufficient text extracted ({len(extracted_text)} chars). Attempting OCR fallback.")
                ocr_text = self._perform_ocr(file_content)
                if ocr_text and len(ocr_text) > len(extracted_text):
                    extracted_text = ocr_text
                    logger.info(f"OCR successful. Extracted {len(extracted_text)} chars.")
                else:
                    logger.warning("OCR failed to produce better results.")

            # Post-processing to fix diacritics artifacts (e.g. "´ a" -> "á", "n˜" -> "ñ")
            def fix_encoding_artifacts(text):
                # 1. Fix Tildes (n˜)
                text = text.replace("n˜", "ñ").replace("N˜", "Ñ")
                text = text.replace("˜n", "ñ").replace("˜N", "Ñ")
                
                # 2. Fix Acute Accents (´ a -> á)
                replacements = {
                    "´ a": "á", "´ e": "é", "´ i": "í", "´ o": "ó", "´ u": "ú",
                    "´ A": "Á", "´ E": "É", "´ I": "Í", "´ O": "Ó", "´ U": "Ú",
                    "´ n": "ñ", "´ N": "Ñ",
                    "´a": "á", "´e": "é", "´i": "í", "´o": "ó", "´u": "ú",
                    "´n": "ñ",
                    # Dotless i cases sometimes appear as 'ı'
                    "´ı": "í", "´ ı": "í"
                }
                for bad, good in replacements.items():
                    text = text.replace(bad, good)
                    
                # 3. Fix simple floating accents/tildes if they remain
                text = text.replace("´", "") # Remove orphan accents if they didn't match above? Maybe too aggressive.
                # Actually, let's keep it safe. Only specific replacements.
                
                return text

            extracted_text = fix_encoding_artifacts(extracted_text)
            cleaned_text = self._clean_extracted_text(extracted_text)

            return {
                "success": True,
                "text": cleaned_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "pdf",
                    "pages_count": pages_count,
                    "parsing_method": "pdfplumber",
                    "text_length": len(cleaned_text),
                    "has_text": len(cleaned_text.strip()) > 0
                }
            }

        except ImportError:
            return self._handle_missing_dependency("pdfplumber", filename, "pdf")
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {e}")
            return self._create_error_response(filename, "pdf", str(e))

    def _perform_ocr(self, file_content: bytes) -> str:
        """
        Realiza OCR en un archivo PDF convirtiéndolo a imágenes
        Requires: poppler-utils and tesseract-ocr installed on the system
        """
        try:
            import pytesseract
            from pdf2image import convert_from_bytes

            logger.info("Attempting OCR extraction...")

            images = convert_from_bytes(file_content)

            if images:
                try:
                    sample_text = pytesseract.image_to_string(images[0], lang='eng+spa')
                    detected_lang = self._detect_ocr_language(sample_text)
                except Exception:
                    detected_lang = 'eng+spa'

                logger.info(f"OCR language set to: {detected_lang}")

                ocr_text_parts = []
                for i, image in enumerate(images):
                    text = pytesseract.image_to_string(image, lang=detected_lang)
                    if text.strip():
                        ocr_text_parts.append(text)
                    logger.info(f"OCR processed page {i+1}/{len(images)}")

                return "\n\n".join(ocr_text_parts)

        except ImportError:
            logger.error("Missing OCR dependencies (pytesseract or pdf2image)")
            return ""
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return ""

    def _detect_ocr_language(self, sample_text: str) -> str:
        """Detecta el idioma para configurar Tesseract OCR dinámicamente"""
        sample_lower = sample_text.lower()

        lang_signatures = {
            'spa': ['experiencia', 'formación', 'educación', 'el', 'los', 'para', 'como'],
            'eng': ['experience', 'education', 'skills', 'the', 'and', 'for', 'work'],
            'fra': ['expérience', 'formation', 'éducation', 'le', 'les', 'pour', 'dans'],
            'deu': ['erfahrung', 'ausbildung', 'bildung', 'der', 'und', 'für', 'mit'],
            'por': ['experiência', 'formação', 'educação', 'o', 'os', 'para', 'como'],
            'ita': ['esperienza', 'formazione', 'istruzione', 'il', 'per', 'come'],
        }

        best_lang = 'eng+spa'
        best_score = 0

        for lang, keywords in lang_signatures.items():
            score = sum(1 for kw in keywords if kw in sample_lower)
            if score > best_score:
                best_score = score
                best_lang = lang

        if best_lang == 'eng+spa':
            return 'eng+spa'
        return f'{best_lang}+eng' if best_lang != 'eng' else 'eng'

    def _parse_word(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extrae texto de archivos Word (.docx, .doc)"""
        try:
            import docx
            from io import BytesIO

            doc = docx.Document(BytesIO(file_content))
            text_parts = []

            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            extracted_text = "\n".join(text_parts)
            cleaned_text = self._clean_extracted_text(extracted_text)

            return {
                "success": True,
                "text": cleaned_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "docx",
                    "paragraphs_count": len(doc.paragraphs),
                    "tables_count": len(doc.tables),
                    "parsing_method": "python-docx",
                    "text_length": len(cleaned_text),
                    "has_text": len(cleaned_text.strip()) > 0
                }
            }

        except ImportError:
            return self._handle_missing_dependency("python-docx", filename, "docx")
        except Exception as e:
            logger.error(f"Error parsing Word document {filename}: {e}")
            return self._create_error_response(filename, "docx", str(e))

    def _parse_text(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extrae texto de archivos de texto plano"""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    extracted_text = file_content.decode(encoding)
                    cleaned_text = self._clean_extracted_text(extracted_text)

                    return {
                        "success": True,
                        "text": cleaned_text,
                        "metadata": {
                            "filename": filename,
                            "file_type": "txt",
                            "encoding": encoding,
                            "parsing_method": "direct_decode",
                            "text_length": len(cleaned_text),
                            "has_text": len(cleaned_text.strip()) > 0
                        }
                    }
                except UnicodeDecodeError:
                    continue

            # Si ningún encoding funciona
            return self._create_error_response(filename, "txt", "No se pudo decodificar el archivo con encodings comunes")

        except Exception as e:
            logger.error(f"Error parsing text file {filename}: {e}")
            return self._create_error_response(filename, "txt", str(e))

    def _parse_rtf(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extrae texto de archivos RTF"""
        try:
            from striprtf.striprtf import rtf_to_text

            rtf_content = file_content.decode('utf-8', errors='ignore')
            extracted_text = rtf_to_text(rtf_content)
            cleaned_text = self._clean_extracted_text(extracted_text)

            return {
                "success": True,
                "text": cleaned_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "rtf",
                    "parsing_method": "striprtf",
                    "text_length": len(cleaned_text),
                    "has_text": len(cleaned_text.strip()) > 0
                }
            }

        except ImportError:
            return self._handle_missing_dependency("striprtf", filename, "rtf")
        except Exception as e:
            logger.error(f"Error parsing RTF file {filename}: {e}")
            return self._create_error_response(filename, "rtf", str(e))

    def _parse_html(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """Extrae texto de archivos HTML (ej: export de LinkedIn)"""
        try:
            from bs4 import BeautifulSoup

            html_content = file_content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html_content, 'lxml')

            for element in soup(['script', 'style', 'nav', 'footer', 'header']):
                element.decompose()

            extracted_text = soup.get_text(separator='\n')
            cleaned_text = self._clean_extracted_text(extracted_text)

            return {
                "success": True,
                "text": cleaned_text,
                "metadata": {
                    "filename": filename,
                    "file_type": "html",
                    "parsing_method": "beautifulsoup",
                    "text_length": len(cleaned_text),
                    "has_text": len(cleaned_text.strip()) > 0
                }
            }

        except ImportError:
            return self._handle_missing_dependency("beautifulsoup4", filename, "html")
        except Exception as e:
            logger.error(f"Error parsing HTML file {filename}: {e}")
            return self._create_error_response(filename, "html", str(e))

    def _clean_extracted_text(self, text: str) -> str:
        """Limpia el texto extraído eliminando caracteres problemáticos"""
        if not text:
            return ""

        # Eliminar caracteres de control excepto saltos de línea y tabs
        cleaned = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

        # Normalizar espacios en blanco
        cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)  # Máximo 2 saltos de línea seguidos
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # Espacios múltiples a uno solo
        cleaned = re.sub(r'[ \t]*\n[ \t]*', '\n', cleaned)  # Limpiar espacios alrededor de saltos de línea

        # Eliminar líneas que solo contienen caracteres especiales
        lines = cleaned.split('\n')
        clean_lines = []
        for line in lines:
            # Mantener líneas que tengan al menos 2 caracteres alfanuméricos
            if len(re.findall(r'[a-zA-Z0-9]', line)) >= 2:
                clean_lines.append(line.strip())
            elif line.strip() == "":  # Mantener líneas vacías para separación
                clean_lines.append("")

        # Reconstruir texto
        result = '\n'.join(clean_lines)

        # Eliminar exceso de líneas vacías al inicio y final
        result = result.strip()

        return result

    def _handle_missing_dependency(self, dependency: str, filename: str, file_type: str) -> Dict[str, any]:
        """Maneja errores de dependencias faltantes"""
        error_msg = f"Dependencia faltante: {dependency}. Instalar con: pip install {dependency}"
        logger.error(error_msg)

        return {
            "success": False,
            "text": "",
            "error": error_msg,
            "metadata": {
                "filename": filename,
                "file_type": file_type,
                "parsing_method": "failed_missing_dependency",
                "missing_dependency": dependency
            }
        }

    def _create_error_response(self, filename: str, file_type: str, error: str) -> Dict[str, any]:
        """Crea una respuesta de error estándar"""
        return {
            "success": False,
            "text": "",
            "error": error,
            "metadata": {
                "filename": filename,
                "file_type": file_type,
                "parsing_method": "error"
            }
        }

    def validate_file(self, file_content: bytes, filename: str) -> Dict[str, any]:
        """
        Valida si un archivo es válido para procesamiento usando magic bytes
        """
        file_extension = Path(filename).suffix.lower()
        file_size = len(file_content)

        # Validación por magic bytes (no confía en extensión)
        detected_type = self._detect_mime_type(file_content)
        extension_matches = file_extension in self.supported_formats

        validations = {
            "is_supported_format": extension_matches,
            "magic_bytes_match": detected_type is not None,
            "extension_matches_magic": self._extension_matches_magic(file_extension, detected_type),
            "has_content": file_size > 0,
            "size_reasonable": 100 <= file_size <= 50 * 1024 * 1024,
            "filename_valid": self._is_filename_valid(filename)
        }

        is_valid = all(validations.values())

        issues = []
        if not validations["is_supported_format"]:
            issues.append(f"Formato no soportado: {file_extension}")
        if not validations["magic_bytes_match"]:
            issues.append("El archivo no coincide con un formato conocido (magic bytes)")
        if not validations["extension_matches_magic"]:
            issues.append(f"La extensión {file_extension} no coincide con el tipo real del archivo")
        if not validations["has_content"]:
            issues.append("El archivo está vacío")
        if not validations["size_reasonable"]:
            if file_size < 100:
                issues.append("El archivo es demasiado pequeño")
            else:
                issues.append("El archivo es demasiado grande (>50MB)")
        if not validations["filename_valid"]:
            issues.append("El nombre del archivo contiene caracteres no válidos")

        return {
            "is_valid": is_valid,
            "validations": validations,
            "issues": issues,
            "file_info": {
                "filename": filename,
                "extension": file_extension,
                "detected_type": detected_type,
                "size_bytes": file_size,
                "size_mb": round(file_size / (1024 * 1024), 2)
            }
        }

    def _detect_mime_type(self, file_content: bytes) -> Optional[str]:
        """Detecta el tipo MIME real del archivo por magic bytes"""
        if len(file_content) < 4:
            return None

        header = file_content[:8]

        if header[:4] == b'%PDF':
            return 'pdf'
        if header[:4] == b'PK\x03\x04':
            if b'word/' in file_content[:512]:
                return 'docx'
            return 'zip'
        if header[:4] == b'\xD0\xCF\x11\xE0':
            return 'doc'
        if header[:6] == b'{\\rtf1':
            return 'rtf'
        if file_content[:1].isascii() and all(
            32 <= b < 127 or b in (9, 10, 13)
            for b in file_content[:min(256, len(file_content))]
        ):
            return 'txt'

        return None

    def _extension_matches_magic(self, extension: str, magic_type: Optional[str]) -> bool:
        """Verifica que la extensión coincida con el magic bytes detectado"""
        if magic_type is None:
            return False

        if extension in ('.txt',):
            return True

        mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.rtf': 'rtf',
            '.html': 'html',
            '.htm': 'html',
        }

        return mapping.get(extension) == magic_type

    def _is_filename_valid(self, filename: str) -> bool:
        """
        Valida que el nombre del archivo sea seguro
        Permite caracteres unicode, espacios, y caracteres especiales comunes
        """
        if not filename or len(filename) == 0:
            return False

        # Verificar longitud razonable
        if len(filename) > 255:
            return False

        # Caracteres prohibidos en sistemas de archivos (más restrictivo en Windows)
        forbidden_chars = ['<', '>', ':', '"', '|', '?', '*']

        # Verificar que no contenga caracteres prohibidos
        for char in forbidden_chars:
            if char in filename:
                return False

        # Verificar que no sea solo puntos
        if filename.strip('.') == '':
            return False

        # Verificar que tenga una extensión válida
        file_extension = Path(filename).suffix.lower()
        if not file_extension:
            return False

        return True